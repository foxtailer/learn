
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import os
import sys


def get_script_directory():
    # Get the full path to the script
    script_path = os.path.realpath(__file__)
    
    # Get the directory containing the script
    script_dir = os.path.dirname(script_path)
    
    return script_dir

def get_weeks(year, month):
    first_day = datetime(year, month, 1)
    if first_day.weekday() == 0:  # if it's already Monday
        start_date = first_day
    else:
        start_date = first_day + timedelta(days=(7 - first_day.weekday()))

    weeks = []
    week_num = 1
    while start_date.month == month:
        end_date = start_date + timedelta(days=4)
        if end_date.month != month:
            end_date = datetime(end_date.year, end_date.month, 1) - timedelta(days=1)
        weeks.append((year, month, week_num, start_date.strftime('%Y-%m-%d'), end_date.strftime('%Y-%m-%d')))
        start_date += timedelta(days=7)
        week_num += 1

    return weeks

def get_weeks_range(start_year, start_month, end_year, end_month):
    result = []
    year, month = start_year, start_month
    while year < end_year or (year == end_year and month <= end_month):
        result.extend(get_weeks(year, month))
        if month == 12:
            month = 1
            year += 1
        else:
            month += 1
    return result

def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), kwargs.get(edge))
            edge_el.set(qn('w:sz'), '4')  # size in eighths of a point
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'auto')
            tcPr.append(edge_el)

def create_document(data):
    script_path = os.path.realpath(__file__)
    script_dir = os.path.dirname(script_path)

    doc = Document()
    table = doc.add_table(rows=1, cols=5)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Year'
    hdr_cells[1].text = 'Month'
    hdr_cells[2].text = 'Week Number'
    hdr_cells[3].text = 'Start Date'
    hdr_cells[4].text = 'End Date'

    for row in data:
        cells = table.add_row().cells
        cells[0].text = str(row[0])
        cells[1].text = str(row[1])
        cells[2].text = str(row[2])
        cells[3].text = row[3]
        cells[4].text = row[4]

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
            set_cell_border(cell, top="single", left="single", bottom="single", right="single")

    doc.save(f"{script_dir}\{start_year}-{start_month} {end_year}-{end_month}.docx")

# Example usage
pattern = r"^\d{4}-\d{2} \d{4}-\d{2}$"
while 1:
    user_input = input("Year-month(start) Year-month(end): ").strip()
    if re.match(pattern, user_input):
        break
    else:
        print("Something wrong. Correct date format: yyyy-mm yyyy-mm"
              "Example: 2024-01 2024-02")

start, end = user_input.split()
start_year, start_month = [int(x) for x in start.split("-")]
end_year, end_month = [int(x) for x in end.split("-")]

data = get_weeks_range(start_year, start_month, end_year, end_month)
create_document(data)
