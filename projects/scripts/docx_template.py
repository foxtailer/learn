from docx import Document
import re

def replace_text_in_paragraph(paragraph, replacements):
    for key, val in replacements.items():
        if key in paragraph.text:
            # Combine all runs into one string
            paragraph_text = paragraph.text
            # Replace the key with the value
            paragraph_text = paragraph_text.replace(key, val)
            
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            
            # Adding new text with preserved formatting
            run = paragraph.add_run()
            run.text = paragraph_text
            # Preserve formatting
            for r in paragraph.runs:
                run.bold = r.bold
                run.italic = r.italic
                run.underline = r.underline
                run.font.size = r.font.size
                run.font.name = r.font.name

def replace_text_in_docx(doc_path, replacements):
    # Load the document
    doc = Document(doc_path)
    
    # Iterate over paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # Iterate over tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    # Save the modified document
    doc.save('modified_' + doc_path)

# Define the path to your .docx file
doc_path = 'Анализ_финансового_состояния_Таибова.docx'

# Define the replacements in the format {'search_text': 'replace_text'}
replacements = {
    '{{some word}}': 'replacement word'
}

# Call the function to replace text
replace_text_in_docx(doc_path, replacements)
